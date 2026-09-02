"""Generate KEYSTONE Project Report as Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading("KEYSTONE — Field Service Management Platform", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Project Report")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(14)

doc.add_paragraph()

# Meta info
meta = [
    ("Project Name:", "KEYSTONE Field Service Management Platform"),
    ("Developer:", "[Your Name]"),
    ("Date:", "August 2026"),
    ("GitHub Repository:", "https://github.com/kalidas1284/keystone"),
    ("Live Application:", "https://keystone-web-vnzs.onrender.com"),
    ("API Documentation:", "https://keystone-api-za1v.onrender.com/swagger-ui.html"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.add_run(label + " ").bold = True
    p.add_run(value)

doc.add_page_break()

# --- Section 1 ---
doc.add_heading("1. Overview", level=1)
doc.add_heading("1.1 Introduction", level=2)
doc.add_paragraph(
    "KEYSTONE is a full-stack Field Service Management (FSM) web application designed for "
    "commercial facilities operations. It enables organizations to manage customers, sites, "
    "technicians, work orders, scheduling, inventory, and SLA compliance from a single operations "
    "console. A dedicated customer self-service portal allows facility contacts to submit and track "
    "service requests without calling the dispatch desk."
)
doc.add_paragraph(
    "The project was built as part of the Zidio Java internship specification, implementing core "
    "FSM features (F1–F9) with a production-ready deployment on Render."
)

doc.add_heading("1.2 Problem Statement", level=2)
doc.add_paragraph("Commercial facility operators need to:")
for item in [
    "Dispatch technicians to multiple customer sites efficiently",
    "Track work order status from creation through completion",
    "Monitor SLA deadlines by job priority",
    "Manage parts inventory and labor time on each job",
    "Give customers visibility into their open requests",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph(
    "KEYSTONE addresses these needs with role-based workflows for admins, managers, dispatchers, "
    "technicians, and customers."
)

doc.add_heading("1.3 Key Features", level=2)
features = [
    ("Authentication & RBAC", "JWT-based login with five roles: Admin, Manager, Dispatcher, Technician, Customer"),
    ("Customer & Site Management", "Customers with multiple building sites; paginated site lists"),
    ("Work Order Lifecycle", "NEW → ASSIGNED → SCHEDULED → IN_PROGRESS → COMPLETED → CLOSED (or CANCELLED)"),
    ("Technician Dispatch", "Assign technicians, schedule jobs, weekly calendar with conflict detection"),
    ("Field Operations", "Technician field board (/field) for mobile-friendly job actions"),
    ("SLA Monitoring", "Priority-based due windows; ON_TRACK, AT_RISK, BREACHED, MET statuses"),
    ("Time & Parts Tracking", "Log labor minutes and inventory parts per work order with roll-up totals"),
    ("Attachments", "Upload photos, PDFs, and documents (up to 10 MB)"),
    ("Inventory", "Stock levels, adjustments, low-stock alerts on dashboard"),
    ("Dashboard & Reports", "Live metrics, charts, technician/customer/site summaries, CSV export"),
    ("Customer Portal", "Submit requests, view status history (no internal cost data exposed)"),
    ("Notifications", "In-app alerts for status changes, scheduling, SLA breaches, portal requests"),
    ("API Documentation", "Swagger UI for all REST endpoints"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Module"
t.rows[0].cells[1].text = "Description"
for mod, desc in features:
    row = t.add_row().cells
    row[0].text = mod
    row[1].text = desc
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Note: ").bold = True
p.add_run("Invoicing and payment processing are intentionally out of scope per the KEYSTONE specification.")

doc.add_heading("1.4 Target Users", level=2)
users = [
    ("Admin", "Full system access, user provisioning, reports"),
    ("Manager", "Operations oversight, close work orders, reports"),
    ("Dispatcher", "Create/assign/schedule work orders, manage customers"),
    ("Technician", "Execute field jobs, log time and parts"),
    ("Customer", "Portal: submit and track service requests"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Role"
t.rows[0].cells[1].text = "Primary Use"
for role, use in users:
    row = t.add_row().cells
    row[0].text = role
    row[1].text = use

doc.add_page_break()

# --- Section 2 ---
doc.add_heading("2. Tech Stack", level=1)

def add_tech_table(doc, heading, items):
    doc.add_heading(heading, level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Technology"
    t.rows[0].cells[1].text = "Purpose"
    for tech, purpose in items:
        row = t.add_row().cells
        row[0].text = tech
        row[1].text = purpose
    doc.add_paragraph()

add_tech_table(doc, "2.1 Frontend", [
    ("React 19", "UI component library"),
    ("TypeScript", "Type-safe JavaScript"),
    ("Vite 8", "Build tool and dev server"),
    ("Tailwind CSS 4", "Utility-first styling"),
    ("React Router 7", "Client-side routing with role guards"),
    ("Axios", "HTTP client for REST API calls"),
    ("Chart.js", "Dashboard charts (status breakdown, workload)"),
    ("React Toastify", "User feedback notifications"),
])
add_tech_table(doc, "2.2 Backend", [
    ("Java 21", "Programming language"),
    ("Spring Boot 4.0.7", "Application framework"),
    ("Spring Security", "Authentication and authorization"),
    ("JWT (JSON Web Tokens)", "Stateless session management"),
    ("Spring Data JPA", "Object-relational mapping"),
    ("PostgreSQL 16", "Relational database"),
    ("Flyway", "Versioned database migrations"),
    ("SpringDoc OpenAPI 3", "Swagger API documentation"),
    ("BCrypt", "Password hashing"),
    ("Maven", "Build and dependency management"),
])
add_tech_table(doc, "2.3 Deployment & DevOps", [
    ("Docker", "Backend containerization"),
    ("Render", "Cloud hosting (API + static frontend + PostgreSQL)"),
    ("GitHub", "Source code repository and version control"),
    ("render.yaml", "Infrastructure-as-code blueprint"),
])

doc.add_heading("2.4 Development Tools", level=2)
for item in [
    "Cursor IDE — Development environment",
    "Git — Version control",
    "JUnit + Mockito — Backend unit and integration tests",
    "H2 Database — In-memory DB for automated tests",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# --- Section 3 ---
doc.add_heading("3. Architecture & Screenshots", level=1)
doc.add_heading("3.1 System Architecture", level=2)
doc.add_paragraph(
    "KEYSTONE follows a classic three-tier architecture:"
)
arch = doc.add_paragraph()
arch.add_run(
    "React SPA (Vite/TS)  →  Spring Boot API (Java 21, JWT)  →  PostgreSQL (keystone_db)\n\n"
    "Request Flow:\n"
    "1. User interacts with the React single-page application\n"
    "2. Frontend sends authenticated HTTP requests to /api/* endpoints\n"
    "3. Spring Security validates JWT and checks role permissions\n"
    "4. Service layer enforces business rules (status transitions, SLA, immutability)\n"
    "5. JPA repositories persist data to PostgreSQL\n"
    "6. DTOs separate API responses from internal entity structure"
).font.name = "Consolas"

doc.add_heading("3.2 Backend Layer Structure", level=2)
layers = doc.add_paragraph()
layers.add_run(
    "controller/  → REST API endpoints\n"
    "service/     → Business logic\n"
    "repository/  → Spring Data JPA interfaces\n"
    "entity/      → JPA domain models\n"
    "dto/         → Request/response objects\n"
    "config/      → Security, Flyway, demo seeder\n"
    "security/    → JWT filter\n"
    "exception/   → Global exception handler"
).font.name = "Consolas"

doc.add_heading("3.3 Database Schema (Key Entities)", level=2)
entities = [
    ("users", "Login accounts with role"),
    ("customers", "Commercial facility clients"),
    ("sites", "Physical locations belonging to a customer"),
    ("technicians", "Field workforce linked to user accounts"),
    ("work_orders", "Core job records with status, priority, SLA due date"),
    ("work_order_status_history", "Audit trail of every status change"),
    ("work_order_time_logs", "Labor minutes logged per job"),
    ("work_order_parts", "Inventory parts consumed per job"),
    ("work_order_attachments", "File uploads per job"),
    ("schedules", "Technician calendar entries"),
    ("inventory_items", "Parts catalog with stock quantities"),
    ("notifications", "In-app user notifications"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Entity"
t.rows[0].cells[1].text = "Description"
for ent, desc in entities:
    row = t.add_row().cells
    row[0].text = ent
    row[1].text = desc

doc.add_heading("3.4 Security Model", level=2)
for item in [
    "Authentication: POST /api/auth/login returns a signed JWT (24-hour expiry)",
    "Authorization: Each endpoint protected by role via @PreAuthorize",
    "Portal isolation: Customer portal excludes internal costs, notes, and technician IDs",
    "Terminal immutability: CLOSED/CANCELLED work orders cannot be edited",
    "CORS: Configured for localhost (dev) and *.onrender.com (production)",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("3.5 SLA Engine", level=2)
sla = [("URGENT", "4 hours"), ("HIGH", "24 hours"), ("MEDIUM", "72 hours"), ("LOW", "7 days")]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Priority"
t.rows[0].cells[1].text = "Due Window"
for pri, window in sla:
    row = t.add_row().cells
    row[0].text = pri
    row[1].text = window
doc.add_paragraph(
    "Statuses: ON_TRACK → AT_RISK (≤25% time remaining) → BREACHED (past due) → MET (completed on time). "
    "A scheduled background task runs every 5 minutes to detect breaches and send notifications."
)

doc.add_heading("3.6 Screenshots", level=2)
doc.add_paragraph(
    "Insert screenshots of the following pages from the live application:"
)
screenshots = [
    ("Figure 1", "Login Page — Keystone login screen with role-based access"),
    ("Figure 2", "Admin Dashboard — Operations dashboard with live metrics, SLA cards, and charts"),
    ("Figure 3", "Work Order Details — Status, time tracking, parts, attachments, status history"),
    ("Figure 4", "Technician Field Board — Mobile-friendly field view for today's jobs"),
    ("Figure 5", "Customer Portal — Self-service portal to submit and track requests"),
    ("Figure 6", "Reports Page — Operational reports with CSV export"),
    ("Figure 7", "Swagger API — OpenAPI/Swagger UI documenting all REST endpoints"),
]
for fig, desc in screenshots:
    p = doc.add_paragraph()
    p.add_run(f"{fig}: ").bold = True
    p.add_run(desc)
    doc.add_paragraph("[ Insert screenshot here ]", style="Intense Quote")

doc.add_page_break()

# --- Section 4 ---
doc.add_heading("4. Deployment", level=1)
doc.add_heading("4.1 Live URLs", level=2)
urls = [
    ("Web Application", "https://keystone-web-vnzs.onrender.com"),
    ("REST API", "https://keystone-api-za1v.onrender.com"),
    ("Swagger UI", "https://keystone-api-za1v.onrender.com/swagger-ui.html"),
    ("Health Check", "https://keystone-api-za1v.onrender.com/api/health"),
    ("GitHub Repository", "https://github.com/kalidas1284/keystone"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Service"
t.rows[0].cells[1].text = "URL"
for svc, url in urls:
    row = t.add_row().cells
    row[0].text = svc
    row[1].text = url

doc.add_heading("4.2 Demo Credentials", level=2)
doc.add_paragraph("Password for all accounts: password123")
creds = [
    ("Admin", "admin@keystone.local"),
    ("Manager", "manager@keystone.local"),
    ("Dispatcher", "dispatcher@keystone.local"),
    ("Technician", "tech@keystone.local"),
    ("Customer", "customer@keystone.local"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Role"
t.rows[0].cells[1].text = "Email"
for role, email in creds:
    row = t.add_row().cells
    row[0].text = role
    row[1].text = email

doc.add_page_break()

# --- Section 5 ---
doc.add_heading("5. Conclusion", level=1)
doc.add_heading("5.1 Summary", level=2)
doc.add_paragraph(
    "KEYSTONE successfully implements a production-grade Field Service Management platform covering "
    "the complete operational workflow: customer onboarding, work order creation, technician dispatch, "
    "field execution, SLA monitoring, inventory management, and customer self-service. The application "
    "demonstrates proficiency in full-stack Java web development with modern frontend practices, secure "
    "REST API design, and cloud deployment."
)

doc.add_heading("5.2 Achievements", level=2)
for item in [
    "All core FSM features (F1–F9) implemented per specification",
    "Five-role RBAC with JWT authentication",
    "SLA engine with automated breach notifications",
    "Customer portal with data isolation (no internal field leakage)",
    "Live deployment on Render with PostgreSQL",
    "Swagger API documentation",
    "Automated backend tests (JUnit)",
    "Flyway database migrations for reproducible schema",
    "CSV report export",
]:
    doc.add_paragraph("✓ " + item, style="List Bullet")

doc.add_heading("5.3 Out of Scope (By Design)", level=2)
for item in [
    "Invoicing and payment processing (per KEYSTONE spec)",
    "Email/SMS external notifications (in-app only)",
    "Mobile native app (responsive web UI used instead)",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("5.4 Future Enhancements", level=2)
for item in [
    "Email/SMS notification integration",
    "Real-time WebSocket updates for dispatch board",
    "GPS-based technician location tracking",
    "Mobile PWA with offline support",
    "Advanced analytics and predictive SLA modeling",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph()
end = doc.add_paragraph("— End of Report —")
end.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = r"c:\intern\keystoone platform\docs\KEYSTONE_Project_Report.docx"
doc.save(out)
print(f"Saved: {out}")
